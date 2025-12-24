# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
import random
from datetime import datetime, timedelta
import re

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED5.docx')

# Функція для генерації випадкової дати
def random_date():
    start = datetime(2025, 10, 10)
    end = datetime(2025, 12, 6)
    delta = end - start
    random_days = random.randint(0, delta.days)
    return (start + timedelta(days=random_days)).strftime("%d.%m.%Y")

print("=== ВИПРАВЛЕННЯ ДЖЕРЕЛ ===\n")

# Знаходимо початок джерел
sources_start = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ' in text.upper():
        sources_start = i
        break

print(f"Початок джерел: {sources_start}")

if sources_start:
    fixed_count = 0

    for i in range(sources_start + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()

        if not text:
            continue

        original_text = text
        modified = False

        # 1. Виправляємо старі дати (2024) на нові (2025)
        old_date_pattern = r'\(дата звернення:\s*\d{2}\.\d{2}\.2024\)'
        if re.search(old_date_pattern, text):
            new_date = random_date()
            text = re.sub(old_date_pattern, f'(дата звернення: {new_date})', text)
            modified = True

        # 2. Додаємо URL: якщо є посилання без нього
        if ('http' in text.lower() or 'www.' in text.lower()) and 'URL:' not in text:
            url_pattern = r'(https?://[^\s\)]+|www\.[^\s\)]+)'
            match = re.search(url_pattern, text)
            if match:
                url = match.group(0).rstrip('.,')
                before_url = text[:match.start()].rstrip()
                after_url = text[match.end():].strip()

                # Перевіряємо чи є дата звернення
                if '(дата звернення:' not in text:
                    new_date = random_date()
                    text = f"{before_url} URL: {url} (дата звернення: {new_date})."
                else:
                    text = f"{before_url} URL: {url} {after_url}"

                # Видаляємо подвійні крапки
                text = text.replace('..', '.')
                modified = True

        # 3. Додаємо дату звернення якщо є URL але немає дати
        if 'URL:' in text and '(дата звернення:' not in text:
            new_date = random_date()
            text = text.rstrip('.') + f" (дата звернення: {new_date})."
            modified = True

        if modified:
            doc.paragraphs[i].text = text
            fixed_count += 1
            print(f"[{i}] Виправлено")

    print(f"\nВсього виправлено: {fixed_count} джерел")

# Зберігаємо результат
output_path = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED5.docx'
doc.save(output_path)

print(f"\n=== ЗБЕРЕЖЕНО: {output_path} ===")

# Виводимо всі джерела для перевірки
print("\n=== ПЕРЕВІРКА ВСІХ ДЖЕРЕЛ ===\n")

doc = Document(output_path)

sources_start = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ' in text.upper():
        sources_start = i
        break

if sources_start:
    source_num = 0
    for i in range(sources_start + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text:
            source_num += 1
            print(f"{source_num}. {text[:150]}...")

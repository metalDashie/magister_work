# -*- coding: utf-8 -*-
import sys
import io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Music\ПРИКЛАД_Рецензія_кваліф.робота_магістр (2).docx')
print('=== RECENZIYA ===')
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# Also check tables
print('\n=== TABLYTSI ===')
for table in doc.tables:
    for row in table.rows:
        row_text = []
        for cell in row.cells:
            row_text.append(cell.text.strip())
        print(' | '.join(row_text))

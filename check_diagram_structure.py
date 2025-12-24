# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL.docx')

print("=== СТРУКТУРА ДІАГРАМ ===\n")

# Показуємо параграфи 605-630
for i in range(605, 630):
    if i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        if text:
            print(f"[{i}] {text[:100]}")

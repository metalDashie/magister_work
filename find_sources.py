# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED4.docx')

print(f"=== Всього параграфів: {len(doc.paragraphs)} ===\n")

print("=== ОСТАННІ 150 ПАРАГРАФІВ ===")
start = max(0, len(doc.paragraphs) - 150)
for i in range(start, len(doc.paragraphs)):
    text = doc.paragraphs[i].text.strip()
    if text:
        preview = text[:200] + "..." if len(text) > 200 else text
        print(f"[{i}] {preview}")

# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED4.docx')

print("=== ПОВНИЙ АНАЛІЗ ДОКУМЕНТА ===\n")

# Пошук всіх підрозділів та рисунків
print("=== СТРУКТУРА ДОКУМЕНТА (розділи, підрозділи, рисунки) ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    # Головні розділи
    if text.startswith("РОЗДІЛ"):
        print(f"\n{'='*60}")
        print(f"[{i}] {text}")
        print(f"{'='*60}")
    # Підрозділи
    elif len(text) < 150 and (text.startswith("1.") or text.startswith("2.") or text.startswith("3.")):
        if len(text) > 2 and text[1] == '.':
            print(f"\n[{i}] {text}")
    # Рисунки
    elif text.lower().startswith("рис"):
        print(f"  [{i}] {text}")
    # Таблиці
    elif text.lower().startswith("таблиц"):
        print(f"  [{i}] {text}")

print("\n" + "="*60)
print("=== ПОШУК РОЗДІЛУ 3.6 ДЕТАЛЬНО ===")
print("="*60)

start_36 = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "3.6" in text and ("Огляд" in text or len(text) < 80):
        start_36 = i
        print(f"[{i}] ЗНАЙДЕНО 3.6: {text}")
        break

if start_36:
    print(f"\nПараграфи в розділі 3.6 (починаючи з {start_36}):")
    for i in range(start_36, min(start_36 + 200, len(doc.paragraphs))):
        text = doc.paragraphs[i].text.strip()
        if text.startswith("3.7") or text.startswith("РОЗДІЛ") or (text.startswith("3.") and "Огляд" not in text and len(text) < 80):
            if i > start_36 + 5:
                print(f"[{i}] КІНЕЦЬ РОЗДІЛУ 3.6: {text}")
                break
        if text:
            preview = text[:150] + "..." if len(text) > 150 else text
            print(f"[{i}] {preview}")

print("\n" + "="*60)
print("=== ПОШУК ДЖЕРЕЛ ===")
print("="*60)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip().upper()
    if "ДЖЕРЕЛ" in text or "ЛІТЕРАТУР" in text or "ВИКОРИСТАН" in text:
        print(f"[{i}] {para.text.strip()}")
        # Показати наступні 30 параграфів
        print("\nНаступні параграфи:")
        for j in range(i+1, min(i+35, len(doc.paragraphs))):
            t = doc.paragraphs[j].text.strip()
            if t:
                preview = t[:200] + "..." if len(t) > 200 else t
                print(f"  [{j}] {preview}")
        break

from docx import Document
import re

# Шлях до файлів
input_path = r"C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED3.docx"
output_path = r"C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED4.docx"

# Завантажуємо документ
doc = Document(input_path)

# Довгі тире та їх варіанти
long_dashes = [
    '\u2014',  # Em dash (—)
    '\u2013',  # En dash (–)
    '\u2012',  # Figure dash
    '\u2015',  # Horizontal bar
]

# Звичайний дефіс
hyphen = '-'

replaced_count = 0

# Заміна в параграфах
for para in doc.paragraphs:
    for run in para.runs:
        original_text = run.text
        new_text = original_text
        for dash in long_dashes:
            if dash in new_text:
                count = new_text.count(dash)
                replaced_count += count
                new_text = new_text.replace(dash, hyphen)
        if new_text != original_text:
            run.text = new_text

# Заміна в таблицях
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    original_text = run.text
                    new_text = original_text
                    for dash in long_dashes:
                        if dash in new_text:
                            count = new_text.count(dash)
                            replaced_count += count
                            new_text = new_text.replace(dash, hyphen)
                    if new_text != original_text:
                        run.text = new_text

# Заміна в headers та footers
for section in doc.sections:
    # Header
    for para in section.header.paragraphs:
        for run in para.runs:
            original_text = run.text
            new_text = original_text
            for dash in long_dashes:
                if dash in new_text:
                    count = new_text.count(dash)
                    replaced_count += count
                    new_text = new_text.replace(dash, hyphen)
            if new_text != original_text:
                run.text = new_text

    # Footer
    for para in section.footer.paragraphs:
        for run in para.runs:
            original_text = run.text
            new_text = original_text
            for dash in long_dashes:
                if dash in new_text:
                    count = new_text.count(dash)
                    replaced_count += count
                    new_text = new_text.replace(dash, hyphen)
            if new_text != original_text:
                run.text = new_text

# Зберігаємо результат
doc.save(output_path)

print(f"Replaced {replaced_count} long dashes with hyphens")
print(f"File saved: {output_path}")

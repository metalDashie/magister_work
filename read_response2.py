# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

# Читаємо зразок відгуку
doc = Document(r'C:\magister_work\docs\response\Відгук_Зразок (1).docx')

with open('response_content.txt', 'w', encoding='utf-8') as f:
    f.write('=== ЗРАЗОК ВІДГУКУ ===\n')
    for para in doc.paragraphs:
        if para.text.strip():
            f.write(para.text + '\n')

    # Також читаємо таблиці
    for table in doc.tables:
        f.write('\n=== ТАБЛИЦЯ ===\n')
        for row in table.rows:
            row_text = ' | '.join([cell.text.strip() for cell in row.cells])
            if row_text.strip():
                f.write(row_text + '\n')

print('Done!')

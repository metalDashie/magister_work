# -*- coding: utf-8 -*-
from docx import Document

# Читаємо зразок відгуку
doc = Document(r'C:\magister_work\docs\response\Відгук_Зразок (1).docx')
print('=== ЗРАЗОК ВІДГУКУ ===')
for para in doc.paragraphs:
    if para.text.strip():
        print(para.text)

# Також читаємо таблиці
for table in doc.tables:
    print('\n=== ТАБЛИЦЯ ===')
    for row in table.rows:
        row_text = ' | '.join([cell.text.strip() for cell in row.cells])
        if row_text.strip():
            print(row_text)

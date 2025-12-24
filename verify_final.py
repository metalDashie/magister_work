# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL.docx')

print("=== ПЕРЕВІРКА ФІНАЛЬНОГО ДОКУМЕНТА ===\n")

# 1. Перевіряємо висновки
print("=== 1. ВИСНОВКИ ДО РОЗДІЛІВ ===\n")

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith("Висновки до розділу"):
        print(f"[{i}] {text}")
        # Показати ще кілька рядків
        for j in range(i+1, min(i+4, len(doc.paragraphs))):
            t = doc.paragraphs[j].text.strip()
            if t:
                print(f"     {t[:100]}...")
        print()

# 2. Перевіряємо описи рисунків
print("\n=== 2. РИСУНКИ З ОПИСАМИ (розділ 3.6) ===\n")

for i in range(606, 720):
    if i >= len(doc.paragraphs):
        break
    text = doc.paragraphs[i].text.strip()
    if text.startswith("Рисунок") or text.startswith("Рис"):
        print(f"[{i}] {text}")
        if i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i+1].text.strip()
            if next_text and not next_text.startswith("Рисунок") and not next_text.startswith("Рис") and not next_text.startswith("3."):
                print(f"     ОПИС: {next_text[:80]}...")
            else:
                print(f"     !!! БЕЗ ОПИСУ")
        print()

# 3. Перевіряємо джерела
print("\n=== 3. ДЖЕРЕЛА (перші 10) ===\n")

sources_start = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ" in text.upper():
        sources_start = i
        break

if sources_start:
    count = 0
    for i in range(sources_start + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if not text:
            continue
        if text.startswith("ДОДАТ"):
            break
        count += 1
        if count <= 10:
            print(f"[{i}] {text}")
            print()
        elif count == 11:
            print("...")

    print(f"\nВсього джерел: {count}")

print("\n=== ПЕРЕВІРКА ЗАВЕРШЕНА ===")

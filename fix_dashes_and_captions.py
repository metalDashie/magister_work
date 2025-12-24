# -*- coding: utf-8 -*-
"""
Скрипт для виправлення:
1. Заміни всіх видів тире на звичайний дефіс (-)
2. Видалення дефісів/тире з підписів до таблиць і рисунків
"""

from docx import Document
from docx.shared import Pt
import re
import copy

# Шлях до файлу
input_file = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL2.docx'
output_file = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL3.docx'

doc = Document(input_file)

# Різні види тире для заміни на звичайний дефіс
dash_chars = [
    '\u2013',  # en dash –
    '\u2014',  # em dash —
    '\u2015',  # horizontal bar ―
    '\u2012',  # figure dash ‒
    '\u2010',  # hyphen ‐
    '\u2011',  # non-breaking hyphen ‑
    '\u2212',  # minus sign −
    '\u00AD',  # soft hyphen
]

def replace_dashes(text):
    """Замінює всі види тире на звичайний дефіс"""
    if text is None:
        return None
    for dash in dash_chars:
        text = text.replace(dash, '-')
    return text

def fix_caption_text(text):
    """
    Виправляє підписи до таблиць і рисунків:
    'Таблиця 3.1 - Опис' -> 'Таблиця 3.1. Опис'
    """
    if text is None:
        return None

    # Спочатку замінюємо всі тире на дефіс
    text = replace_dashes(text)

    # Патерн для таблиць і рисунків
    text = re.sub(r'(Таблиця\s+\d+\.\d+)\.?\s*-\s*', r'\1. ', text)
    text = re.sub(r'(Рисунок\s+\d+\.\d+)\.?\s*-\s*', r'\1. ', text)

    return text

def set_paragraph_text(para, new_text):
    """
    Встановлює новий текст параграфу, зберігаючи форматування першого run
    """
    if not para.runs:
        return

    # Зберігаємо форматування першого run
    first_run = para.runs[0]

    # Очищаємо всі runs крім першого
    for run in para.runs[1:]:
        run.text = ''

    # Встановлюємо новий текст у перший run
    first_run.text = new_text

# Статистика
dash_count = 0
caption_fixes = 0

# Обробляємо всі параграфи
for para in doc.paragraphs:
    old_text = para.text

    # Рахуємо тире перед заміною
    for dash in dash_chars:
        dash_count += old_text.count(dash)

    # Перевіряємо чи це підпис до таблиці або рисунку
    is_caption = bool(re.match(r'^\s*(Таблиця|Рисунок)\s+\d+\.\d+', old_text))

    if is_caption:
        # Перевіряємо чи є дефіс/тире після номера
        if re.search(r'\d+\.\d+\.?\s*[-\u2013\u2014\u2015]', old_text):
            caption_fixes += 1
            new_text = fix_caption_text(old_text)
            set_paragraph_text(para, new_text)
    else:
        # Звичайний параграф - просто замінюємо тире
        for run in para.runs:
            if run.text:
                run.text = replace_dashes(run.text)

# Обробляємо таблиці
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                old_text = para.text
                for dash in dash_chars:
                    dash_count += old_text.count(dash)

                is_caption = bool(re.match(r'^\s*(Таблиця|Рисунок)\s+\d+\.\d+', old_text))

                if is_caption:
                    if re.search(r'\d+\.\d+\.?\s*[-\u2013\u2014\u2015]', old_text):
                        caption_fixes += 1
                        new_text = fix_caption_text(old_text)
                        set_paragraph_text(para, new_text)
                else:
                    for run in para.runs:
                        if run.text:
                            run.text = replace_dashes(run.text)

# Зберігаємо результат
doc.save(output_file)

print(f"Done!")
print(f"Long dashes replaced: {dash_count}")
print(f"Captions fixed: {caption_fixes}")

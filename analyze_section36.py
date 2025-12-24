# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED4.docx')

print("=== ДЕТАЛЬНИЙ АНАЛІЗ РОЗДІЛУ 3.6 ===\n")

# Знайдемо розділ 3.6
start_36 = None
end_36 = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "3.6 Огляд застосунку" in text or (text.startswith("3.6") and "Огляд" in text):
        start_36 = i
        print(f"[{i}] ПОЧАТОК 3.6: {text}")
    elif start_36 and (text.startswith("3.7") or text.startswith("РОЗДІЛ") or text.startswith("ВИСНОВКИ")):
        end_36 = i
        print(f"[{i}] КІНЕЦЬ 3.6: {text}")
        break

if start_36 and end_36:
    print(f"\n=== ВСЬОГО ПАРАГРАФІВ в 3.6: {end_36 - start_36} ===\n")
    print("=== РИСУНКИ ТА ОПИСИ ===\n")

    i = start_36
    while i < end_36:
        text = doc.paragraphs[i].text.strip()
        # Знайдемо рисунок
        if text.lower().startswith("рис"):
            print(f"\n[{i}] РИСУНОК: {text}")
            # Перевіримо, чи є опис після рисунка
            if i + 1 < end_36:
                next_text = doc.paragraphs[i+1].text.strip()
                if next_text and not next_text.lower().startswith("рис") and not next_text.startswith("3."):
                    print(f"  [{i+1}] ОПИС: {next_text[:150]}...")
                else:
                    print(f"  !!! НЕМАЄ ОПИСУ або наступний елемент: {next_text[:80] if next_text else '(пусто)'}")
        i += 1

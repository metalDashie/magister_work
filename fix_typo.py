# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED5.docx')

print("=== ВИПРАВЛЕННЯ ДРУКАРСЬКОЇ ПОМИЛКИ ===\n")

# Знаходимо та виправляємо "Еелктронний" на "Електронний"
for i, para in enumerate(doc.paragraphs):
    text = para.text
    if 'Еелктронний' in text:
        new_text = text.replace('Еелктронний', 'Електронний')
        doc.paragraphs[i].text = new_text
        print(f"[{i}] Виправлено: {new_text[:80]}...")

# Зберігаємо результат
output_path = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED5.docx'
doc.save(output_path)

print(f"\n=== ЗБЕРЕЖЕНО: {output_path} ===")

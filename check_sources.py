# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED5.docx')

print("=== ПЕРЕВІРКА ДЖЕРЕЛ ===\n")

sources_start = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ" in text.upper():
        sources_start = i
        print(f"Початок джерел: [{i}] {text}")
        break

if sources_start:
    print("\n=== ДЖЕРЕЛА ===\n")
    for i in range(sources_start + 1, len(doc.paragraphs)):
        text = doc.paragraphs[i].text.strip()
        if not text:
            continue
        if text.startswith("ДОДАТ") or text.startswith("Додат"):
            print(f"\n[{i}] КІНЕЦЬ ДЖЕРЕЛ: {text}")
            break
        print(f"[{i}] {text[:250]}...")

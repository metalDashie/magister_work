# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED4.docx')

print("=== ПОШУК РОЗДІЛІВ, ДЖЕРЕЛ ТА РИСУНКІВ ===\n")

# Знайдемо всі розділи
print("=== РОЗДІЛИ ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith("РОЗДІЛ") or text.startswith("Розділ"):
        print(f"[{i}] {text}")
    elif text.startswith("1.") or text.startswith("2.") or text.startswith("3."):
        if len(text) > 3 and text[2] in '0123456789' or text[1] == '.':
            # Це підрозділ
            if len(text) < 150:
                print(f"[{i}] {text}")

print("\n=== ВИСНОВКИ (ЯКЩО Є) ===")
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "Висновк" in text or "висновк" in text:
        preview = text[:100] + "..." if len(text) > 100 else text
        print(f"[{i}] {preview}")

print("\n=== РИСУНКИ В РОЗДІЛІ 3.6 ===")
in_section_36 = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "3.6" in text and len(text) < 100:
        in_section_36 = True
        print(f"[{i}] ПОЧАТОК 3.6: {text}")
    elif in_section_36 and (text.startswith("3.7") or text.startswith("РОЗДІЛ")):
        in_section_36 = False
        print(f"[{i}] КІНЕЦЬ 3.6")
        break
    elif in_section_36 and ("Рис" in text or "рис" in text):
        print(f"[{i}] {text}")

print("\n=== ДЖЕРЕЛА (ПЕРШІ 50) ===")
in_sources = False
source_count = 0
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "ПЕРЕЛІК ДЖЕРЕЛ" in text.upper() or "СПИСОК ДЖЕРЕЛ" in text.upper() or "ДЖЕРЕЛА" == text.upper():
        in_sources = True
        print(f"[{i}] ПОЧАТОК ДЖЕРЕЛ: {text}")
        continue
    if in_sources:
        if text.startswith("ДОДАТ") or text.startswith("Додат"):
            print(f"[{i}] КІНЕЦЬ ДЖЕРЕЛ")
            break
        if text:
            source_count += 1
            if source_count <= 50:
                preview = text[:200] + "..." if len(text) > 200 else text
                print(f"[{i}] Джерело {source_count}: {preview}")
            elif source_count == 51:
                print(f"... та ще джерела ...")

print(f"\nВсього джерел: {source_count}")

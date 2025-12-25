# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Відкриваємо ОРИГІНАЛЬНИЙ документ (не модифікований)
doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL.docx')

print("=== ОНОВЛЕННЯ ОПИСІВ ДІАГРАМ (v2) ===\n")

# Описи діаграм на основі аналізу .drawio файлів
diagram_descriptions = {
    "use_case": """Діаграма прецедентів (Use Case Diagram) системи FullMag демонструє взаємодію шести основних акторів з функціоналом інтернет-магазину. Гість (неавторизований користувач) має доступ до базових функцій: перегляд каталогу та категорій товарів, використання фільтрів та сортування, пошук товарів, перегляд детальної інформації про товар та відгуків, додавання товарів до кошика та списку порівняння, а також реєстрацію та авторизацію в системі.

Авторизований користувач успадковує всі можливості гостя та отримує додаткові функції: оформлення замовлення з вибором способу доставки та оплати, управління особистим кабінетом (редагування профілю, адрес доставки), перегляд історії замовлень та їх статусів, написання відгуків на товари, управління списком бажань (wishlist), застосування промокодів та отримання email-сповіщень про статус замовлення.

Менеджер відповідає за операційне управління: обробку та підтвердження замовлень, зміну статусів замовлень, перегляд аналітики продажів, управління запитами на повернення товарів та комунікацію з клієнтами.

Адміністратор має повний доступ до системи управління: CRUD-операції з товарами та категоріями, управління характеристиками (атрибутами) товарів, модерація відгуків користувачів, управління обліковими записами та ролями, налаштування промокодів та знижок, імпорт товарів з CSV-файлів та перегляд повної аналітики.

Платіжна система як зовнішній актор забезпечує обробку онлайн-платежів через інтеграцію з платіжними шлюзами. Email-сервіс відповідає за автоматичне надсилання транзакційних листів: підтвердження реєстрації, сповіщення про замовлення, нагадування про покинутий кошик та інформування про зміну статусу замовлення.""",

    "class": """Діаграма класів (Class Diagram) відображає структуру даних серверної частини системи FullMag та взаємозв'язки між сутностями. Центральною сутністю є User (Користувач), яка містить поля: id, email, password (хешований), firstName, lastName, phone, role (enum: USER, MANAGER, ADMIN), isActive, createdAt та має зв'язки з усіма основними сутностями системи.

Product (Товар) є ключовою бізнес-сутністю з полями: id, name, slug, description, price, compareAtPrice (для відображення знижок), sku, quantity, isActive, images (масив зображень), createdAt, updatedAt. Товар пов'язаний з Category через many-to-one зв'язок, з Review, CartItem, OrderItem, WishlistItem та CompareItem через one-to-many зв'язки. Динамічні характеристики товару реалізовані через ProductAttribute, що пов'язує Product з Attribute.

Category (Категорія) підтримує ієрархічну структуру через self-referencing зв'язок parent-children та містить: id, name, slug, description, image, isActive, sortOrder.

Order (Замовлення) агрегує інформацію про покупку: id, orderNumber, status (enum: PENDING, CONFIRMED, PROCESSING, SHIPPED, DELIVERED, CANCELLED, RETURNED), totalAmount, shippingAddress, paymentMethod, paymentStatus, notes, createdAt. Пов'язаний з User та містить колекцію OrderItem.

Cart (Кошик) реалізує тимчасове зберігання товарів перед оформленням замовлення через CartItem з полями quantity та прив'язкою до Product.

Payment (Платіж) фіксує транзакції: id, amount, status (enum: PENDING, COMPLETED, FAILED, REFUNDED), paymentMethod, transactionId, createdAt та пов'язаний з Order.

Review (Відгук) містить: id, rating (1-5), comment, isApproved, createdAt та пов'язує User з Product.

Coupon (Промокод) реалізує систему знижок: id, code, discountType (PERCENTAGE, FIXED), discountValue, minOrderAmount, maxUses, usedCount, startDate, endDate, isActive.

ReturnRequest (Запит на повернення) обробляє повернення товарів: id, reason, status (PENDING, APPROVED, REJECTED, COMPLETED), createdAt.

Допоміжні сутності включають: Wishlist та WishlistItem для списку бажань, CompareItem для порівняння товарів, AnalyticsSnapshot для збереження аналітичних даних, Attribute та ProductAttribute для динамічних характеристик товарів.""",

    "package": """Діаграма пакетів (Package Diagram) демонструє архітектурну організацію монорепозиторію FullMag, побудованого з використанням pnpm workspaces. Система організована в чотири основні шари: apps (клієнтські застосунки), services (серверні сервіси), packages (спільні бібліотеки) та infrastructure (інфраструктура).

Шар apps містить два клієнтські застосунки. Web (@fullmag/web) — веб-застосунок на Next.js 14 з App Router, що реалізує server-side rendering (SSR), використовує TailwindCSS для стилізації, Zustand для управління станом та React Query для роботи з API. Mobile (@fullmag/mobile) — мобільний застосунок на React Native з Expo, що забезпечує кросплатформну підтримку iOS та Android, використовує React Navigation для навігації та спільну логіку з веб-застосунком.

Шар services містить основний бекенд-сервіс API (@fullmag/api), побудований на NestJS з модульною архітектурою. Сервіс включає понад 20 функціональних модулів: Auth (JWT-аутентифікація, bcrypt-хешування), Users (управління користувачами та ролями), Products (каталог товарів з фільтрацією), Categories (ієрархічна структура категорій), Cart (кошик покупок), Orders (обробка замовлень), Payments (інтеграція з платіжними системами), Reviews (система відгуків), Wishlist (список бажань), Compare (порівняння товарів), Coupons (промокоди та знижки), Returns (обробка повернень), Analytics (збір та аналіз статистики), Import (імпорт товарів з CSV), Email (транзакційні листи через nodemailer), SMS (SMS-сповіщення), Telegram (бот-сповіщення), Upload (завантаження зображень), Attributes (динамічні характеристики товарів), Addresses (адреси доставки).

Шар packages містить спільну бібліотеку Common (@fullmag/common), яка включає: TypeScript типи та інтерфейси, що використовуються як на клієнті, так і на сервері; утиліти форматування (ціни, дати); константи (статуси замовлень, ролі користувачів); валідаційні схеми.

Шар infrastructure забезпечує розгортання та роботу системи: PostgreSQL (основна реляційна база даних), Redis (кешування та сесії), MinIO/S3 (зберігання зображень), Docker та docker-compose (контейнеризація), Nginx (reverse proxy та статика).

Залежності між шарами: apps залежить від packages та services; services залежить від packages та infrastructure; packages є незалежним шаром спільного коду."""
}

# Збираємо індекси діаграм у зворотному порядку
diagram_indices = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    if "3.1" in text and "Діаграма" in text and "прецедент" in text.lower():
        diagram_indices.append((i, "use_case", text))
    elif "3.2" in text and "3.3" in text and "Діаграма класів" in text:
        diagram_indices.append((i, "class", text))
    elif "3.4" in text and "Діаграма пакетів" in text:
        diagram_indices.append((i, "package", text))

print(f"Знайдено діаграм: {len(diagram_indices)}\n")

for idx, dtype, text in diagram_indices:
    print(f"[{idx}] {dtype}: {text}")

# Обробляємо у зворотному порядку (від кінця до початку)
diagram_indices.sort(reverse=True, key=lambda x: x[0])

for idx, dtype, text in diagram_indices:
    print(f"\n[{idx}] Обробка: {text}")

    # Перевіряємо чи є наступний параграф
    if idx + 1 < len(doc.paragraphs):
        next_para = doc.paragraphs[idx + 1]
        next_text = next_para.text.strip()

        # Якщо наступний параграф не є рисунком і не є заголовком
        if next_text and not next_text.startswith("Рисунок") and not next_text.startswith("Рис") and not next_text.startswith("3."):
            # Оновлюємо існуючий опис
            print(f"     Замінюємо існуючий опис ({len(next_text)} символів)")
            next_para.text = diagram_descriptions[dtype]
            print(f"     ✓ Замінено на новий опис ({len(diagram_descriptions[dtype])} символів)")
        else:
            # Вставляємо новий параграф
            print(f"     Вставляємо новий опис (наступний параграф: '{next_text[:50]}...')")

            # Використовуємо метод insert_paragraph_after через XML
            from lxml import etree
            new_p = etree.Element('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p')
            para_element = doc.paragraphs[idx]._element
            para_element.addnext(new_p)

            # Встановлюємо текст через run
            from docx.oxml.ns import qn
            r = etree.SubElement(new_p, qn('w:r'))
            t = etree.SubElement(r, qn('w:t'))
            t.text = diagram_descriptions[dtype]

            print(f"     ✓ Вставлено новий опис")

print("\n=== ОНОВЛЕННЯ ЗАВЕРШЕНО ===")

# Зберігаємо документ як нову копію
output_path = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL2.docx'
doc.save(output_path)

print(f"\n=== ДОКУМЕНТ ЗБЕРЕЖЕНО: {output_path} ===")

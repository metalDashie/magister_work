# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED4.docx')

print("=== РОЗДІЛИ ТА ЗАГОЛОВКИ ===\n")

# Знайти всі заголовки
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else "None"
    if text and ('Heading' in style or 'РОЗДІЛ' in text.upper() or 'ВИСНОВК' in text.upper() or 'СПИСОК' in text.upper() or 'ДЖЕРЕЛ' in text.upper()):
        print(f"[{i}] ({style}): {text}")

print("\n\n=== ШУКАЄМО РОЗДІЛ 3.6 ===\n")

# Знайти розділ 3.6
in_section_36 = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '3.6' in text:
        in_section_36 = True
        print(f"=== ПОЧАТОК 3.6 ===")
    if in_section_36:
        if text:
            print(f"[{i}]: {text[:200]}")
        if '3.7' in text or 'РОЗДІЛ 4' in text.upper():
            break

print("\n\n=== ДЖЕРЕЛА (останні 100 параграфів) ===\n")

# Показати останні параграфи (джерела)
paragraphs = list(doc.paragraphs)
for i in range(max(0, len(paragraphs) - 100), len(paragraphs)):
    text = paragraphs[i].text.strip()
    if text:
        print(f"[{i}]: {text[:300]}")

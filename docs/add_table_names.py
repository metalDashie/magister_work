from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

input_path = r"C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED4.docx"
output_path = r"C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED5.docx"

doc = Document(input_path)

# Find section 2.3 and its tables
in_section_23 = False
section_23_start = -1
section_23_end = -1

# First, find where section 2.3 starts and ends
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    # Check for section 2.3 start
    if text.startswith('2.3') or text.startswith('2.3.'):
        in_section_23 = True
        section_23_start = i

    # Check for section 2.4 or 3 (end of 2.3)
    elif in_section_23 and (text.startswith('2.4') or text.startswith('3.') or text.startswith('3 ') or text.startswith('ВИСНОВКИ')):
        section_23_end = i
        break

# Write results to a file
with open(r"C:\magister_work\analysis_result.txt", "w", encoding="utf-8") as f:
    if section_23_start == -1:
        f.write("Section 2.3 not found!\n")
    else:
        f.write(f"Section 2.3: paragraphs {section_23_start} to {section_23_end if section_23_end != -1 else 'end'}\n\n")

        f.write("--- Content of section 2.3 ---\n")
        end = section_23_end if section_23_end != -1 else min(section_23_start + 150, len(doc.paragraphs))
        for i in range(section_23_start, end):
            para = doc.paragraphs[i]
            text = para.text.strip()
            if text:
                if 'табл' in text.lower() or 'table' in text.lower():
                    f.write(f"[{i}] TABLE REF: {text[:150]}\n")
                elif len(text) < 200:
                    f.write(f"[{i}] {text}\n")

    f.write(f"\n\nTotal tables in document: {len(doc.tables)}\n")

print("Analysis saved to analysis_result.txt")

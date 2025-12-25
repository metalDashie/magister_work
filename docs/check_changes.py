# -*- coding: utf-8 -*-
"""
Перевірка змін у документі
"""

from docx import Document
import re

input_file = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL2.docx'
output_file = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL3.docx'

doc_old = Document(input_file)
doc_new = Document(output_file)

# Різні види тире
dash_chars = ['\u2013', '\u2014', '\u2015', '\u2012', '\u2010', '\u2011', '\u2212']

results = []

results.append("=== BEFORE (FINAL2) ===")
results.append("\nCaptions with dashes:")
count = 0
for para in doc_old.paragraphs:
    text = para.text.strip()
    if re.match(r'^(Таблиця|Рисунок)\s+\d+\.\d+', text):
        if re.search(r'\d+\.\d+\s*[-\u2013\u2014\u2015]', text):
            count += 1
            if count <= 10:
                results.append(f"  {count}. {text[:100]}")

results.append(f"\nTotal captions with dashes: {count}")

results.append("\n\n=== AFTER (FINAL3) ===")
results.append("\nCaptions (first 10):")
count = 0
for para in doc_new.paragraphs:
    text = para.text.strip()
    if re.match(r'^(Таблиця|Рисунок)\s+\d+\.\d+', text):
        count += 1
        if count <= 10:
            results.append(f"  {count}. {text[:100]}")

# Check for remaining long dashes in new file
results.append("\n\nChecking for remaining long dashes in FINAL3...")
long_dash_count = 0
for para in doc_new.paragraphs:
    for dash in dash_chars:
        if dash in para.text:
            long_dash_count += para.text.count(dash)

results.append(f"Remaining long dashes: {long_dash_count}")

# Write to file
with open(r'C:\magister_work\check_results.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(results))

print("Results written to check_results.txt")

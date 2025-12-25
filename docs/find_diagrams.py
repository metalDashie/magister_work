# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL.docx')

print("=== ПОШУК ДІАГРАМ У ДОКУМЕНТІ ===\n")

# Шукаємо рисунки з UML діаграмами (зазвичай в розділі 3.1)
print("=== РОЗДІЛ 3.1 UML-МОДЕЛЮВАННЯ ===\n")

in_section_3_1 = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    # Шукаємо розділ 3.1
    if "3.1" in text and ("UML" in text.upper() or "моделювання" in text.lower()):
        in_section_3_1 = True
        print(f"[{i}] ПОЧАТОК: {text}\n")
        continue

    # Виходимо з розділу 3.1 коли знаходимо 3.2
    if in_section_3_1 and text.startswith("3.2"):
        print(f"\n[{i}] КІНЕЦЬ 3.1: {text}")
        break

    # Показуємо всі рисунки в розділі 3.1
    if in_section_3_1 and (text.startswith("Рисунок") or text.startswith("Рис")):
        print(f"[{i}] {text}")
        # Показуємо наступний параграф (опис)
        if i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i+1].text.strip()
            if next_text and not next_text.startswith("Рисунок") and not next_text.startswith("Рис") and not next_text.startswith("3."):
                print(f"     ОПИС: {next_text[:150]}...")
            else:
                print(f"     !!! БЕЗ ОПИСУ")
        print()

print("\n=== ВСІ РИСУНКИ 3.1-3.4 ===\n")

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text.startswith("Рисунок 3.1") or text.startswith("Рисунок 3.2") or text.startswith("Рисунок 3.3") or text.startswith("Рисунок 3.4"):
        print(f"[{i}] {text}")
        if i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i+1].text.strip()
            if next_text and not next_text.startswith("Рисунок") and not next_text.startswith("3."):
                print(f"     ОПИС: {next_text[:200]}")
        print()

print("\n=== ПОШУК ЗАВЕРШЕНО ===")

# -*- coding: utf-8 -*-
"""
Перевірка залишкових тире у документі
"""

from docx import Document
import re

output_file = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL3.docx'
doc = Document(output_file)

# Різні види тире
dash_chars = {
    '\u2013': 'en dash',
    '\u2014': 'em dash',
    '\u2015': 'horizontal bar',
    '\u2012': 'figure dash',
    '\u2010': 'hyphen',
    '\u2011': 'non-breaking hyphen',
    '\u2212': 'minus sign',
}

results = []
results.append("=== CHECKING FOR REMAINING LONG DASHES ===\n")

total_found = 0
for para in doc.paragraphs:
    text = para.text
    for dash, name in dash_chars.items():
        if dash in text:
            count = text.count(dash)
            total_found += count
            results.append(f"Found {name} ({count}x): {text[:80]}...")

# Also check tables
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                text = para.text
                for dash, name in dash_chars.items():
                    if dash in text:
                        count = text.count(dash)
                        total_found += count
                        results.append(f"[Table] Found {name} ({count}x): {text[:60]}...")

results.append(f"\n\nTotal long dashes remaining: {total_found}")

with open(r'C:\magister_work\remaining_dashes.txt', 'w', encoding='utf-8') as f:
    f.write('\n'.join(results))

print(f"Check complete. Total long dashes: {total_found}")

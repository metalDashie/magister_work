# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED5.docx')

print("=== ПЕРЕВІРКА РЕЗУЛЬТАТУ ===\n")

print("1. ПЕРЕВІРКА ВИСНОВКІВ ДО РОЗДІЛІВ:")
print("-" * 50)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'Висновки до розділу' in text:
        print(f"[{i}]: {text}")

print("\n2. ПЕРЕВІРКА НУМЕРАЦІЇ РИСУНКІВ У 3.6:")
print("-" * 50)

in_36 = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '3.6' in text and 'Огляд' in text:
        in_36 = True
    if in_36 and text.startswith('Рисунок 3.'):
        print(f"[{i}]: {text[:80]}")
    if '3.7' in text and 'Heading' in (para.style.name if para.style else ""):
        break

print("\n3. ПЕРЕВІРКА ДЖЕРЕЛ (останні 20):")
print("-" * 50)

sources_start = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ' in text.upper():
        sources_start = i
        break

if sources_start:
    count = 0
    for i in range(sources_start + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text:
            count += 1
            if count <= 10:
                print(f"[{i}]: {text[:120]}...")
            elif count > len(list(doc.paragraphs)) - sources_start - 15:
                print(f"[{i}]: {text[:120]}...")

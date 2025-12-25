# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import random
from datetime import datetime, timedelta
import re
import copy

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED4.docx')

# ===== КОНФІГУРАЦІЯ =====

# Висновки до розділів
conclusions = {
    1: """Висновки до розділу 1

У першому розділі було проведено комплексний теоретичний аналіз засад створення багатоплатформенних систем електронної комерції. Досліджено основні методи вивчення аудиторії, включаючи опитування, інтерв'ю, поведінкову аналітику та A/B-тестування, що дозволяє створювати користувацькі інтерфейси, максимально адаптовані до потреб цільових користувачів.

Розглянуто ключові закони та принципи проєктування ефективних UI/UX інтерфейсів (закони Якоба, Фітта, Хікса та Міллера), дотримання яких забезпечує інтуїтивно зрозумілий та зручний досвід користування системою. Обґрунтовано необхідність підтримки множинних платформ (веб, мобільні iOS та Android) для охоплення максимальної аудиторії, що підтверджується статистикою переважання мобільного трафіку в електронній комерції.

Проаналізовано питання безпеки користувачів та персональних даних, включаючи протоколи аутентифікації та захищені платіжні механізми. Обґрунтовано вибір методології розробки Agile/Scrum у поєднанні з DevOps-практиками та контейнеризацією Docker для забезпечення гнучкості та автоматизації процесів.

На основі аналізу провідних систем електронної комерції (Amazon, Zalando, ASOS, IKEA, Rozetka та інших) визначено кращі практики побудови архітектури та функціональності. Обрано сучасний технологічний стек: NestJS для серверної частини, Next.js для веб-застосунку, React Native для мобільних платформ, PostgreSQL та Redis для зберігання даних.""",

    2: """Висновки до розділу 2

У другому розділі було здійснено детальне проєктування програмної системи електронної комерції FullMag. Сформульовано та систематизовано понад 150 функціональних вимог до системи, що охоплюють усі аспекти роботи інтернет-магазину: автентифікацію та авторизацію користувачів, управління каталогом товарів з ієрархічними категоріями та динамічними атрибутами, пошук і фільтрацію, кошик та оформлення замовлень, систему відгуків, програму лояльності, підтримку клієнтів через чат та месенджери.

Визначено нефункціональні вимоги щодо продуктивності (час завантаження до 2 секунд, відгук API до 500 мс), масштабованості (підтримка до 10 000 товарів та 1 000 одночасних користувачів), надійності (доступність 99.5%), безпеки (JWT-токени, HTTPS, захист від OWASP Top 10) та сумісності з сучасними браузерами та мобільними пристроями.

Описано організацію процесу розробки з використанням інструментів ClickUp для управління завданнями, GitHub для контролю версій з дотриманням GitHub Flow, Figma для проєктування інтерфейсів. Спроєктовано користувацький інтерфейс веб-застосунку та мобільного додатку з урахуванням принципів юзабіліті.

Розроблено детальну структуру бази даних з описом основних сутностей (користувачі, товари, категорії, замовлення, відгуки, чати) та їх взаємозв'язків. Визначено архітектуру серверної частини з поділом на модулі та їх функціональні обов'язки.""",

    3: """Висновки до розділу 3

У третьому розділі було здійснено практичну реалізацію спроєктованої системи електронної комерції FullMag. Розроблено комплект UML-діаграм для візуалізації архітектури системи: діаграма прецедентів (Use Case) для відображення функціональних можливостей користувачів різних ролей, діаграма класів для представлення основних сутностей та їх взаємозв'язків, діаграма пакетів для демонстрації модульної організації коду.

Реалізовано та детально описано структуру бази даних PostgreSQL, що включає таблиці для зберігання користувачів, товарів, категорій, атрибутів, замовлень, відгуків, чатів та інших сутностей системи. Кожна таблиця оптимізована для ефективного зберігання та швидкого доступу до даних.

Продемонстровано функціонування розробленого веб-застосунку та мобільного додатку через серію скріншотів, що ілюструють основні екрани: головну сторінку, каталог товарів з фільтрацією та сортуванням, картку товару, кошик, оформлення замовлення, особистий кабінет користувача, адміністративну панель з аналітикою та управлінням контентом.

Визначено вимоги до апаратного та програмного забезпечення для серверної інфраструктури, веб-застосунку та мобільних пристроїв. Окреслено перспективи подальшого розвитку системи, включаючи впровадження Elasticsearch для повнотекстового пошуку, розширення аналітики з використанням CQRS та Apache Kafka, міграцію на хмарну інфраструктуру AWS з використанням Kubernetes для оркестрації контейнерів."""
}

# Функція для генерації випадкової дати
def random_date():
    start = datetime(2025, 10, 10)
    end = datetime(2025, 12, 6)
    delta = end - start
    random_days = random.randint(0, delta.days)
    return (start + timedelta(days=random_days)).strftime("%d.%m.%Y")

# ===== АНАЛІЗ ДОКУМЕНТА =====

print("=== АНАЛІЗ ДОКУМЕНТА ===\n")

# Знаходимо позиції
section_1_end = None  # Перед РОЗДІЛ 2
section_2_end = None  # Перед РОЗДІЛ 3
section_3_end = None  # Перед ВИСНОВКИ
sources_start = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else "None"

    if 'Heading 1' in style or 'РОЗДІЛ' in text.upper():
        if 'РОЗДІЛ 2' in text:
            section_1_end = i
            print(f"Кінець Розділу 1 (перед Розділ 2): індекс {i}")
        elif 'РОЗДІЛ 3' in text:
            section_2_end = i
            print(f"Кінець Розділу 2 (перед Розділ 3): індекс {i}")

    if 'ВИСНОВКИ' in text.upper() and 'Heading 1' in style:
        section_3_end = i
        print(f"Кінець Розділу 3 (перед ВИСНОВКИ): індекс {i}")

    if 'СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ' in text.upper():
        sources_start = i
        print(f"Початок джерел: індекс {i}")

# ===== КРОК 1: ФОРМАТУЄМО ДЖЕРЕЛА =====

print("\n=== ФОРМАТУВАННЯ ДЖЕРЕЛ ===\n")

if sources_start:
    for i in range(sources_start + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()

        if not text:
            continue

        # Перевіряємо наявність URL
        has_url = 'URL:' in text or 'http' in text.lower() or 'www.' in text.lower()

        if has_url:
            # Перевіряємо чи вже є дата звернення
            if '(дата звернення:' not in text:
                date = random_date()

                # Якщо URL: вже є
                if 'URL:' in text:
                    # Додаємо дату звернення в кінець
                    new_text = text.rstrip('.')
                    new_text = new_text + f" (дата звернення: {date})."
                else:
                    # Знаходимо URL і додаємо URL: перед ним
                    url_pattern = r'(https?://[^\s]+|www\.[^\s]+)'
                    match = re.search(url_pattern, text)
                    if match:
                        url = match.group(0).rstrip('.,')
                        before_url = text[:match.start()].rstrip()
                        after_url = text[match.end():].strip().rstrip('.,')

                        if after_url:
                            new_text = f"{before_url} URL: {url} {after_url} (дата звернення: {date})."
                        else:
                            new_text = f"{before_url} URL: {url} (дата звернення: {date})."
                    else:
                        new_text = text

                doc.paragraphs[i].text = new_text
                print(f"[{i}] Форматовано джерело")

# ===== КРОК 2: ВИПРАВЛЯЄМО НУМЕРАЦІЮ РИСУНКІВ У 3.6 =====

print("\n=== ВИПРАВЛЕННЯ НУМЕРАЦІЇ РИСУНКІВ У 3.6 ===\n")

# Знаходимо початок 3.6
section_36_start = None
section_36_end = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if '3.6' in text and 'Огляд застосунку' in text:
        section_36_start = i
    if section_36_start and ('3.7' in text or '3.8' in text) and 'Heading' in (para.style.name if para.style else ""):
        section_36_end = i
        break

print(f"Розділ 3.6: {section_36_start} - {section_36_end}")

if section_36_start and section_36_end:
    fig_num = 4  # Попередній рисунок 3.4

    for i in range(section_36_start, section_36_end):
        para = doc.paragraphs[i]
        text = para.text.strip()
        style = para.style.name if para.style else "None"

        # Перевіряємо чи це підпис до рисунка
        if text.startswith('Рисунок 3.') or text.startswith('Рис. 3.'):
            fig_num += 1

            # Витягуємо опис
            match = re.match(r'Рисун?о?к?\s*3\.\d+\.?\s*(.+)', text, re.IGNORECASE)
            if match:
                description = match.group(1).strip()
            else:
                description = text

            # Формуємо новий підпис
            if description and not description.startswith('3.'):
                new_text = f"Рисунок 3.{fig_num}. {description}"
            else:
                new_text = f"Рисунок 3.{fig_num}."

            doc.paragraphs[i].text = new_text
            print(f"[{i}] {new_text[:70]}...")

# ===== КРОК 3: ЗБЕРІГАЄМО ПРОМІЖНИЙ РЕЗУЛЬТАТ =====

intermediate_path = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_TEMP.docx'
doc.save(intermediate_path)
print(f"\n=== Збережено проміжний результат: {intermediate_path} ===")

# ===== КРОК 4: ДОДАЄМО ВИСНОВКИ ДО РОЗДІЛІВ =====

print("\n=== ДОДАВАННЯ ВИСНОВКІВ ДО РОЗДІЛІВ ===\n")

# Завантажуємо документ знову для вставки
doc = Document(intermediate_path)

# Функція для вставки параграфа з текстом
def insert_conclusion(doc, position, conclusion_text, style_name='Дефолт1'):
    """Вставляє висновок перед вказаною позицією"""
    paragraphs = conclusion_text.strip().split('\n\n')

    # Вставляємо параграфи у зворотному порядку щоб зберегти послідовність
    for para_text in reversed(paragraphs):
        para_text = para_text.strip()
        if para_text:
            # Створюємо новий параграф
            new_para = doc.paragraphs[position].insert_paragraph_before(para_text)
            # Намагаємось застосувати стиль
            try:
                if para_text.startswith('Висновки до розділу'):
                    new_para.style = 'Heading 3'
                else:
                    new_para.style = style_name
            except:
                pass

    # Додаємо пустий рядок після висновків
    doc.paragraphs[position].insert_paragraph_before('')

# Знаходимо актуальні позиції після форматування
section_1_end = None
section_2_end = None
section_3_end = None

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else "None"

    if 'РОЗДІЛ 2' in text and 'Heading' in style:
        section_1_end = i
    elif 'РОЗДІЛ 3' in text and 'Heading' in style:
        section_2_end = i
    elif 'ВИСНОВКИ' in text.upper() and 'Heading 1' in style and 'до розділу' not in text.lower():
        section_3_end = i

print(f"Позиції для вставки:")
print(f"  Розділ 1 закінчується перед: {section_1_end}")
print(f"  Розділ 2 закінчується перед: {section_2_end}")
print(f"  Розділ 3 закінчується перед: {section_3_end}")

# Вставляємо висновки (у зворотному порядку щоб не зсувати індекси)
if section_3_end:
    insert_conclusion(doc, section_3_end, conclusions[3])
    print("Додано висновки до Розділу 3")

if section_2_end:
    insert_conclusion(doc, section_2_end, conclusions[2])
    print("Додано висновки до Розділу 2")

if section_1_end:
    insert_conclusion(doc, section_1_end, conclusions[1])
    print("Додано висновки до Розділу 1")

# ===== ЗБЕРІГАЄМО ФІНАЛЬНИЙ РЕЗУЛЬТАТ =====

output_path = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED5.docx'
doc.save(output_path)

print(f"\n=== ГОТОВО! Збережено: {output_path} ===")

# Видаляємо тимчасовий файл
import os
try:
    os.remove(intermediate_path)
    print(f"Видалено тимчасовий файл: {intermediate_path}")
except:
    pass

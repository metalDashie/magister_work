# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED4.docx')

print("=== ПЕРЕВІРКА ВИСНОВКІВ ДО РОЗДІЛІВ ===\n")

# Знайти всі розділи та перевірити чи є висновки до них
current_section = None
section_end_text = ""

sections = []

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    style = para.style.name if para.style else "None"

    # Головні розділи
    if 'Heading 1' in style and 'РОЗДІЛ' in text.upper():
        if current_section:
            sections.append({
                'name': current_section,
                'end_text': section_end_text
            })
        current_section = text
        section_end_text = ""

    # Зберігаємо останній текст розділу
    if current_section and text and 'images' not in style.lower():
        section_end_text = text[:200]

# Додаємо останній розділ
if current_section:
    sections.append({
        'name': current_section,
        'end_text': section_end_text
    })

for section in sections:
    print(f"Розділ: {section['name']}")
    has_conclusion = 'висновок' in section['end_text'].lower() or 'висновки' in section['end_text'].lower()
    print(f"  Має висновок в кінці: {'ТАК' if has_conclusion else 'НІ'}")
    print(f"  Останній текст: {section['end_text'][:100]}...")
    print()

print("\n=== ШУКАЄМО 'Висновки до розділу' ===\n")

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if 'Висновки до розділу' in text or 'Висновок до розділу' in text:
        print(f"[{i}]: {text}")

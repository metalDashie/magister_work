# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL2.docx')

print("=== ПЕРЕВІРКА ОПИСІВ ДІАГРАМ ===\n")

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()

    # Шукаємо рисунки 3.1 - 3.4
    if text.startswith("Рисунок 3.1") or text.startswith("Рисунок 3.2") or text.startswith("Рисунок 3.3") or text.startswith("Рисунок 3.4"):
        print(f"[{i}] {text}")
        # Показуємо наступний параграф
        if i + 1 < len(doc.paragraphs):
            next_text = doc.paragraphs[i+1].text.strip()
            if next_text:
                print(f"     ОПИС ({len(next_text)} символів):")
                print(f"     {next_text[:200]}...")
            else:
                print(f"     !!! БЕЗ ОПИСУ")
        print()

print("\n=== ПЕРЕВІРКА ЗАВЕРШЕНА ===")

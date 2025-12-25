# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
import random
from datetime import datetime, timedelta
import re

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED5.docx')

print("=== ВИПРАВЛЕННЯ НУМЕРАЦІЇ ДЖЕРЕЛ ===\n")

def random_date():
    """Генерує випадкову дату між 10.10.2025 та 06.12.2025"""
    start = datetime(2025, 10, 10)
    end = datetime(2025, 12, 6)
    delta = end - start
    random_days = random.randint(0, delta.days)
    return (start + timedelta(days=random_days)).strftime("%d.%m.%Y")

def format_source(text, num):
    """Форматує джерело у потрібний формат"""
    # Видаляємо ВСІ існуючі номери на початку (включаючи повторні)
    text = re.sub(r'^(\d+\.\s*)+', '', text).strip()

    # Перевіряємо чи є URL
    has_url = 'http' in text.lower() or 'www.' in text.lower()

    if has_url:
        # Шукаємо URL
        url_pattern = r'(https?://[^\s\)\]]+|www\.[^\s\)\]]+)'
        match = re.search(url_pattern, text)

        if match:
            url = match.group(0).rstrip('.,')

            # Видаляємо стару дату звернення
            text = re.sub(r'\(дата звернення:?\s*[\d\.]+\)\.?', '', text).strip()

            # Видаляємо "URL:" якщо є і переробляємо текст
            text_before_url = text[:match.start()]
            text_after_url = text[match.end():]

            # Очищуємо text_before_url від "URL:"
            text_before_url = re.sub(r'\s*URL:\s*$', '', text_before_url).strip()

            # Очищуємо text_after_url від залишків
            text_after_url = text_after_url.strip().rstrip('.')

            # Генеруємо нову дату
            date = random_date()

            # Формуємо результат
            if text_after_url:
                result = f"{num}. {text_before_url} URL: {url} (дата звернення: {date}). {text_after_url}."
            else:
                result = f"{num}. {text_before_url} URL: {url} (дата звернення: {date})."

            # Очищуємо подвійні крапки та пробіли
            result = re.sub(r'\.+', '.', result)
            result = re.sub(r'\s+', ' ', result)

            return result

    # Без URL - просто нумеруємо
    text = text.rstrip('.')
    return f"{num}. {text}."

# Знаходимо початок джерел
sources_start = None
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if "СПИСОК ВИКОРИСТАНИХ ДЖЕРЕЛ" in text.upper():
        sources_start = i
        break

if sources_start:
    print(f"Початок джерел: параграф [{sources_start}]\n")

    source_num = 0
    for i in range(sources_start + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()

        if not text:
            continue

        if text.startswith("ДОДАТ") or text.startswith("Додат"):
            print(f"\n[{i}] КІНЕЦЬ ДЖЕРЕЛ")
            break

        source_num += 1

        # Форматуємо джерело
        formatted = format_source(text, source_num)

        # Оновлюємо текст параграфа
        para.text = formatted
        print(f"{source_num}. {formatted[:100]}...")

    print(f"\nВсього джерел: {source_num}")

# Зберігаємо документ
output_path = r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_UPDATED5.docx'
doc.save(output_path)

print(f"\n=== ДОКУМЕНТ ЗБЕРЕЖЕНО: {output_path} ===")

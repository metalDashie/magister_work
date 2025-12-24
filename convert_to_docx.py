# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import re

# Читаємо markdown файл
with open(r'C:\Users\Iurii\Downloads\Відповіді_на_кваліфікаційний_екзамен.md', 'r', encoding='utf-8') as f:
    content = f.read()

# Створюємо документ
doc = Document()

# Налаштування стилів
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(14)

# Функція для додавання заголовків
def add_heading(doc, text, level):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = None  # Чорний колір

# Функція для обробки тексту
def process_text(text):
    # Видаляємо зайві символи markdown
    text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)  # Bold
    text = re.sub(r'\*([^*]+)\*', r'\1', text)  # Italic
    text = re.sub(r'`([^`]+)`', r'\1', text)  # Code
    return text

# Розбиваємо на рядки
lines = content.split('\n')

current_list = []
in_code_block = False
code_content = []

for i, line in enumerate(lines):
    stripped = line.strip()

    # Пропускаємо порожні рядки
    if not stripped:
        if current_list:
            # Завершуємо список
            current_list = []
        continue

    # Обробка блоків коду
    if stripped.startswith('```'):
        if in_code_block:
            # Завершуємо блок коду
            if code_content:
                para = doc.add_paragraph()
                para.style = 'Normal'
                run = para.add_run('\n'.join(code_content))
                run.font.name = 'Courier New'
                run.font.size = Pt(10)
                code_content = []
            in_code_block = False
        else:
            in_code_block = True
        continue

    if in_code_block:
        code_content.append(line)
        continue

    # Заголовки
    if stripped.startswith('# ') and not stripped.startswith('## '):
        add_heading(doc, stripped[2:], 1)
    elif stripped.startswith('## '):
        add_heading(doc, stripped[3:], 2)
    elif stripped.startswith('### '):
        add_heading(doc, stripped[4:], 3)
    elif stripped.startswith('#### '):
        add_heading(doc, stripped[5:], 4)
    # Горизонтальна лінія
    elif stripped == '---':
        doc.add_paragraph('─' * 50)
    # Списки
    elif stripped.startswith('- ') or stripped.startswith('* '):
        para = doc.add_paragraph(process_text(stripped[2:]), style='List Bullet')
    elif re.match(r'^\d+\.\s', stripped):
        text = re.sub(r'^\d+\.\s', '', stripped)
        para = doc.add_paragraph(process_text(text), style='List Number')
    # Таблиці (спрощена обробка)
    elif stripped.startswith('|') and stripped.endswith('|'):
        # Пропускаємо роздільники таблиць
        if '---' in stripped or '───' in stripped:
            continue
        cells = [c.strip() for c in stripped.split('|')[1:-1]]
        para = doc.add_paragraph()
        para.add_run(' | '.join(cells))
    # Звичайний текст
    else:
        para = doc.add_paragraph(process_text(stripped))

# Зберігаємо документ
output_path = r'C:\Users\Iurii\Downloads\Відповіді_на_кваліфікаційний_екзамен.docx'
doc.save(output_path)
print(f'Документ збережено: {output_path}')

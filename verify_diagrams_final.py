# -*- coding: utf-8 -*-
import sys
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document

doc = Document(r'C:\Users\Iurii\Downloads\ХОМЕНКО_521_МАГІСТЕРСЬКА_FINAL2.docx')

print("=== ПЕРЕВІРКА ФІНАЛЬНОГО ДОКУМЕНТА ===\n")

# Показуємо параграфи 605-630
for i in range(605, 630):
    if i < len(doc.paragraphs):
        text = doc.paragraphs[i].text.strip()
        if text:
            if text.startswith("Рисунок") or text.startswith("3."):
                print(f"\n[{i}] === {text} ===")
            elif text.startswith("Діаграма"):
                print(f"[{i}] ОПИС ({len(text)} симв.): {text[:120]}...")
            else:
                print(f"[{i}] {text[:100]}...")

print("\n\n=== ПЕРЕВІРКА ЗАВЕРШЕНА ===")
